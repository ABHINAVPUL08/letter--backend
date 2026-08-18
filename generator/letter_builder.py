"""Build the Unable to Reach letter as a Word document matching the office sample."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from generator.content import HEADER_OFFICES, LETTERS

NAVY = RGBColor(0x1F, 0x4E, 0x79)


def _set_run_font(
    run: Run,
    font_name: str,
    size_pt: float,
    *,
    complex_script: bool = False,
    bidi_lang: str = "en-US",
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    half_points = str(int(size_pt * 2))
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), half_points)
    sz_cs = rPr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        rPr.append(sz_cs)
    sz_cs.set(qn("w:val"), half_points)
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "en-US")
    if complex_script:
        lang.set(qn("w:bidi"), bidi_lang)


def _set_paragraph_spacing(
    paragraph: Paragraph,
    *,
    before: int = 0,
    after: int = 0,
    line: int | None = None,
    line_rule: str = "auto",
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pPr = paragraph._p.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), line_rule)


def _set_cell_borderless(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_table_borders_none(table: Table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        tblPr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _add_bottom_border(paragraph: Paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_top_border(paragraph: Paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)
    pPr.append(pBdr)


def _write_header(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(title, before=0, after=0, line=240)
    run = title.add_run("LAW OFFICE OF JASPREET SINGH, ESQ")
    _set_run_font(run, "Calibri", 20)
    run.bold = True
    run.font.color.rgb = NAVY

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(tag, before=0, after=6, line=240)
    run = tag.add_run("Attorneys At Law")
    _set_run_font(run, "Times New Roman", 13)
    run.italic = True

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders_none(table)
    table.autofit = True
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    for idx, office in enumerate(HEADER_OFFICES):
        cell = table.rows[0].cells[idx]
        _set_cell_borderless(cell)
        cell.text = ""
        align = align_map.get(office["align"], WD_ALIGN_PARAGRAPH.CENTER)
        for i, line in enumerate(office["lines"]):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = align
            _set_paragraph_spacing(p, before=0, after=0, line=200)
            r = p.add_run(line)
            _set_run_font(r, "Times New Roman", 9)

    line1 = doc.add_paragraph()
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(line1, before=8, after=0, line=200)
    _add_bottom_border(line1)

    admitted = doc.add_paragraph()
    admitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(admitted, before=2, after=2, line=200)
    run = admitted.add_run("***Admitted in NY & CA")
    _set_run_font(run, "Times New Roman", 9)

    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(line2, before=0, after=12, line=200)
    _add_top_border(line2)


def _meta_line(
    doc: Document,
    label: str,
    value: str,
    font: str,
    *,
    complex_script: bool,
    bidi_lang: str,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, before=0, after=2, line=276)
    r1 = p.add_run(f"{label} ")
    _set_run_font(r1, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)
    r1.bold = True
    r2 = p.add_run(value)
    _set_run_font(r2, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.6)


def _write_letter(
    doc: Document,
    *,
    language: str,
    client_name: str,
    a_number_formatted: str,
    letter_date: str,
) -> None:
    lang_key = "english" if language == "english" else language
    tpl = deepcopy(LETTERS[lang_key])
    font = tpl["font"]
    complex_script = bool(tpl["complex_script"])
    bidi_lang = str(tpl.get("bidi_lang") or "en-US")

    _write_header(doc)
    _meta_line(
        doc, tpl["date_label"], letter_date, font, complex_script=complex_script, bidi_lang=bidi_lang
    )
    _meta_line(
        doc, tpl["name_label"], client_name, font, complex_script=complex_script, bidi_lang=bidi_lang
    )
    _meta_line(
        doc,
        tpl["a_number_label"],
        a_number_formatted,
        font,
        complex_script=complex_script,
        bidi_lang=bidi_lang,
    )

    subject = doc.add_paragraph()
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(subject, before=10, after=12, line=276)
    sr = subject.add_run(tpl["subject"])
    _set_run_font(sr, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)
    sr.bold = True

    sal = doc.add_paragraph()
    sal.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(sal, before=0, after=8, line=276)
    r = sal.add_run(tpl["salutation"])
    _set_run_font(r, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)

    for spans in tpl["paragraphs"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(p, before=0, after=10, line=276)
        for span in spans:
            run = p.add_run(span["text"])
            _set_run_font(run, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)
            run.bold = bool(span.get("bold"))

    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(close, before=18, after=6, line=276)
    r = close.add_run(tpl["closing"])
    _set_run_font(r, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)

    sig_line = doc.add_paragraph()
    sig_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(sig_line, before=0, after=4, line=200)
    sr_line = sig_line.add_run("____________________")
    _set_run_font(sr_line, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)

    for i, line in enumerate(
        (tpl["signatory_name"], tpl["signatory_firm"], *tpl["signatory_office"])
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_spacing(p, before=0, after=0, line=240)
        r = p.add_run(line)
        _set_run_font(r, font, 12, complex_script=complex_script, bidi_lang=bidi_lang)
        if i == 0:
            r.bold = True


def build_letter_docx(
    *,
    language: str,
    client_name: str,
    a_number_formatted: str,
    letter_date: str,
    addr_code: str,
    dest_path: Path,
) -> Path:
    doc = Document()
    _configure_page(doc)
    style = doc.styles["Normal"]
    tpl = LETTERS["english" if language == "english" else language]
    style.font.name = tpl["font"]
    style.font.size = Pt(12)
    _write_letter(
        doc,
        language=language,
        client_name=client_name,
        a_number_formatted=a_number_formatted,
        letter_date=letter_date,
    )
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path


def build_combined_letter_docx(
    *,
    native_language: str,
    english_name: str,
    native_name: str,
    a_number_formatted: str,
    native_a_number: str,
    letter_date: str,
    native_date: str,
    dest_path: Path,
) -> Path:
    """English letter on page 1, translated letter on page 2, one Word file."""
    doc = Document()
    _configure_page(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    _write_letter(
        doc,
        language="english",
        client_name=english_name,
        a_number_formatted=a_number_formatted,
        letter_date=letter_date,
    )
    br = doc.add_paragraph()
    _set_paragraph_spacing(br, before=0, after=0, line=240)
    br.add_run().add_break(WD_BREAK.PAGE)
    _write_letter(
        doc,
        language=native_language,
        client_name=native_name,
        a_number_formatted=native_a_number,
        letter_date=native_date,
    )

    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path
